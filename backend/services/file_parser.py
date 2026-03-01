"""
services/file_parser.py
───────────────────────
Extracts plain text from uploaded files.

Supported formats:
  .txt / .md / .csv  — decoded directly (UTF-8 → latin-1 → cp1252 fallback)
  .pdf               — pdfplumber first (better layout), PyPDF2 as fallback
  .docx              — python-docx (paragraphs + table cells)
  .doc               — unsupported; raises a clear error message

All parsing functions are synchronous. FastAPI's async file reading
(await file.read()) is handled in the router before calling extract_text().
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger("neuro.parser")


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(filename: str, content: bytes) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Args:
        filename: Original filename (used only for extension detection).
        content:  Raw file bytes.

    Returns:
        Extracted plain text as a str.

    Raises:
        ValueError: On unsupported type, empty result, or parse failure.
                    Message is safe to surface directly to the user.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md", ".csv"):
        return _parse_text(content)

    if suffix == ".pdf":
        return _parse_pdf(content)

    if suffix == ".docx":
        return _parse_docx(content)

    if suffix == ".doc":
        raise ValueError(
            "Old-format .doc files are not supported. "
            "Please open in Word, save as .docx, and re-upload."
        )

    raise ValueError(
        f"Unsupported file type '{suffix}'. "
        "Accepted formats: .pdf, .docx, .txt, .md"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Plain text / Markdown / CSV
# ─────────────────────────────────────────────────────────────────────────────

def _parse_text(content: bytes) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            text = content.decode(encoding)
            if text.strip():
                return text
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(
        "Could not decode the text file. "
        "Please ensure it is saved as UTF-8 and try again."
    )


# ─────────────────────────────────────────────────────────────────────────────
# PDF  (pdfplumber → PyPDF2 fallback)
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pdf(content: bytes) -> str:
    # Try pdfplumber first — better at preserving layout and handling columns
    text = _pdf_pdfplumber(content)

    if not text.strip():
        logger.info("pdfplumber returned empty — trying PyPDF2 fallback")
        text = _pdf_pypdf2(content)

    if not text.strip():
        raise ValueError(
            "No readable text found in this PDF. "
            "It may be a scanned image. Try the Vision Scan module instead."
        )

    return text


def _pdf_pdfplumber(content: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed — skipping")
        return ""

    try:
        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages.append(page_text.strip())
        return "\n\n".join(pages)
    except Exception as exc:
        logger.warning(f"pdfplumber failed: {exc}")
        return ""


def _pdf_pypdf2(content: bytes) -> str:
    try:
        import PyPDF2
    except ImportError:
        raise ValueError(
            "No PDF parser is installed. "
            "Run: pip install pdfplumber PyPDF2"
        )

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as exc:
        raise ValueError(f"PDF parse error: {exc}")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "python-docx is not installed. "
            "Run: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(content))
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Table cells — joined with pipe separators per row
        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n\n".join(parts)

        if not text.strip():
            raise ValueError("The DOCX file appears to be empty or contains no readable text.")

        return text

    except ValueError:
        raise  # re-raise our own clear messages
    except Exception as exc:
        raise ValueError(f"DOCX parse error: {exc}")